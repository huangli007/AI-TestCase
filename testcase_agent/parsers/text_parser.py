"""文本类文件解析:txt / md / csv / json / docx / pdf。"""

from __future__ import annotations

import csv
import io
import json
import logging
import os

logger = logging.getLogger(__name__)

TEXT_EXTS = {".txt", ".md", ".markdown", ".csv", ".json", ".log", ".yaml", ".yml", ".html", ".htm"}


def extract_text(path: str) -> str:
    """根据扩展名提取文本内容。"""
    ext = os.path.splitext(path)[1].lower()
    if ext in TEXT_EXTS:
        return _read_plain(path, ext)
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".pdf":
        return _read_pdf(path)
    raise ValueError(f"不支持的文件类型: {ext}(文本解析器仅支持 txt/md/csv/json/docx/pdf)")


def _read_plain(path: str, ext: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    if ext == ".csv":
        try:
            rows = list(csv.reader(io.StringIO(content)))
            table = "\n".join(" | ".join(cell.strip() for cell in row) for row in rows if row)
            return f"[CSV 数据表格]\n{table}"
        except Exception:  # noqa: BLE001
            return content
    if ext == ".json":
        try:
            return json.dumps(json.loads(content), ensure_ascii=False, indent=2)
        except Exception:  # noqa: BLE001
            return content
    return content


def _read_docx(path: str) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "解析 .docx 需要安装 python-docx。\n"
            "请运行: pip install python-docx>=1.1.0\n"
            "(或使用项目 venv: 见 README 启动方式)"
        ) from e
    document = docx.Document(path)
    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "解析 .pdf 需要安装 pypdf。\n"
            "请运行: pip install pypdf>=4.2.0"
        ) from e
    reader = PdfReader(path)
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text)
    return "\n\n--- 第 %d 页 ---\n\n".join(pages) if pages else "(PDF 未提取到文本,可能为扫描件,建议以图片方式分析)"
